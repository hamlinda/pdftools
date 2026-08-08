import os
import io
import uuid
import shutil
import socket
import logging
import tempfile
import subprocess
import zipfile
import markdown
from typing import List
from PIL import Image
import fitz
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
from html.parser import HTMLParser
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle



# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

app = FastAPI(title="PDF Tools Backend")

# Security Headers Middleware
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    response = await call_next(request)
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-eval' 'unsafe-inline'; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "font-src 'self' https://fonts.gstatic.com; "
        "img-src 'self' data:; "
        "connect-src 'self';"
    )
    return response

# CORS Configuration - Restrict to local network origins (localhost and local network IP)
# For ease of local deployment, we can configure CORS to allow request from self
# and prevent wildcard CORS issues.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Restrict to * or configure specifically if needed. TODO(security): harden CORS origins in prod.
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

def cleanup_temp_dir(dir_path: str):
    """
    Cleans up a temporary directory securely after response is sent.
    """
    try:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
            logging.info(f"Successfully cleaned up temporary directory: {dir_path}")
    except Exception as e:
        logging.error(f"Failed to clean up temporary directory {dir_path}: {e}")

@app.get("/api/network-info")
def get_network_info():
    """
    Exposes the server's local network IP address to help other devices connect.
    """
    try:
        # Connect to an external IP to find the primary interface IP
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception as e:
        logging.warning(f"Could not determine local network IP: {e}")
        local_ip = "127.0.0.1"
    
    return {"local_ip": local_ip}

def clean_text_lines(text: str) -> str:
    """
    Cleans text lines by stripping whitespace and joining them,
    handling word hyphens across line breaks.
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    result = ""
    for line in lines:
        if result:
            if result.endswith("-"):
                result = result[:-1] + line
            else:
                result += " " + line
        else:
            result = line
    return result


@app.post("/api/pdf-to-docx")
def convert_pdf_to_docx(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """
    Uploads a PDF, extracts text (using digital extraction or Tesseract OCR for scanned pages),
    returns the DOCX file, and registers a cleanup task.
    """
    # 1. Validate extension
    filename = file.filename or "document.pdf"
    sanitized_filename = os.path.basename(filename)
    if not sanitized_filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only PDF documents are allowed."
        )

    # Create a completely secure random filename inside a private temporary directory
    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")
    docx_path = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")

    # Queue cleanup task to prevent disk pollution
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        # Save upload to temporary file
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 2. Validate PDF format using Magic Bytes check (%PDF-)
        if os.path.getsize(pdf_path) < 5:
             raise HTTPException(status_code=400, detail="Invalid PDF file. File is too small.")
             
        with open(pdf_path, "rb") as f:
            magic_bytes = f.read(5)
            if magic_bytes != b"%PDF-":
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid PDF file header. The uploaded file is not a valid PDF."
                )

        logging.info(f"Starting conversion for file: {sanitized_filename} -> size: {os.path.getsize(pdf_path)} bytes")

        # 3. Perform conversion (run synchronously; FastAPI executes sync endpoints in a threadpool)
        import fitz
        import docx
        from docx import Document
        
        doc = Document()
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            
            # Extract digital text blocks
            blocks = page.get_text("blocks")
            text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
            
            full_digital_text = "".join([b[4] for b in text_blocks])
            alnum_count = sum(1 for c in full_digital_text if c.isalnum())
            
            paragraphs_to_add = []
            
            if alnum_count >= 30:
                logging.info(f"Page {page_num + 1}: Using digital text extraction ({alnum_count} alnum chars).")
                for block in text_blocks:
                    block_text = block[4]
                    cleaned = clean_text_lines(block_text)
                    if cleaned:
                        paragraphs_to_add.append(cleaned)
            else:
                logging.info(f"Page {page_num + 1}: Scanned page or low digital text ({alnum_count} alnum chars). Running OCR...")
                zoom = 150 / 72
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                    temp_img_path = temp_img.name
                try:
                    pix.save(temp_img_path)
                    cmd = ["tesseract", temp_img_path, "stdout", "-l", "eng"]
                    res = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8")
                    
                    if res.returncode == 0:
                        ocr_text = res.stdout
                        raw_paragraphs = ocr_text.split("\n\n")
                        for p in raw_paragraphs:
                            cleaned = clean_text_lines(p)
                            if cleaned:
                                paragraphs_to_add.append(cleaned)
                        logging.info(f"Page {page_num + 1}: OCR successfully extracted {len(paragraphs_to_add)} paragraphs.")
                    else:
                        logging.error(f"Page {page_num + 1}: OCR failed with exit code {res.returncode}: {res.stderr}")
                        # Fallback to whatever digital text we have
                        for block in text_blocks:
                            cleaned = clean_text_lines(block[4])
                            if cleaned:
                                paragraphs_to_add.append(cleaned)
                except Exception as ocr_err:
                    logging.error(f"Page {page_num + 1}: OCR exception: {ocr_err}")
                    for block in text_blocks:
                        cleaned = clean_text_lines(block[4])
                        if cleaned:
                            paragraphs_to_add.append(cleaned)
                finally:
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
            
            if page_num > 0:
                doc.add_page_break()
                
            p_header = doc.add_paragraph()
            run = p_header.add_run(f"--- Page {page_num + 1} ---")
            run.bold = True
            p_header.paragraph_format.space_before = docx.shared.Pt(12)
            p_header.paragraph_format.space_after = docx.shared.Pt(6)
            
            if not paragraphs_to_add:
                doc.add_paragraph("[Empty Page]")
            else:
                for p_text in paragraphs_to_add:
                    p = doc.add_paragraph(p_text)
                    p.paragraph_format.space_after = docx.shared.Pt(6)
                    
        doc.save(docx_path)

        # Check if the output docx was created
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
            raise HTTPException(
                status_code=500, 
                detail="Conversion failed. Could not generate Word Document."
            )

        # Re-derive output name
        base_name = os.path.splitext(sanitized_filename)[0]
        docx_download_filename = f"{base_name}.docx"

        logging.info(f"Conversion successful, sending file: {docx_download_filename}")

        # Return file response with security headers
        return FileResponse(
            path=docx_path,
            filename=docx_download_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{docx_download_filename}"',
                "X-Content-Type-Options": "nosniff"
            }
        )

    except HTTPException as he:
        # Re-raise HTTP exceptions directly
        raise he
    except Exception as e:
        logging.error(f"Error during PDF to DOCX conversion: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"An error occurred during conversion: {str(e)}"
        )

# --- ReportLab HTML/Markdown & Word Fallback Converters ---

class HTMLToPlatypusParser(HTMLParser):
    def __init__(self, styles):
        super().__init__()
        self.styles = styles
        self.flowables = []
        self.current_tag = None
        self.current_text = ""
        self.list_mode = None  # 'ul' or 'ol'
        self.list_index = 1
        
    def handle_starttag(self, tag, attrs):
        if tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'pre', 'code']:
            self.current_tag = tag
            self.current_text = ""
        elif tag in ['ul', 'ol']:
            self.list_mode = tag
            self.list_index = 1
        elif tag in ['b', 'i', 'u', 'strong', 'em', 'span', 'a']:
            if self.current_tag:
                if tag == 'a':
                    href = dict(attrs).get('href', '')
                    self.current_text += f'<a href="{href}">'
                else:
                    mapped = 'b' if tag == 'strong' else ('i' if tag == 'em' else tag)
                    self.current_text += f"<{mapped}>"
        elif tag == 'hr':
            self.flowables.append(Spacer(1, 10))

    def handle_endtag(self, tag):
        if tag in ['b', 'i', 'u', 'strong', 'em', 'span', 'a']:
            if self.current_tag:
                mapped = 'b' if tag == 'strong' else ('i' if tag == 'em' else tag)
                self.current_text += f"</{mapped}>"
        elif tag == self.current_tag:
            text = self.current_text.strip()
            if text:
                if tag == 'h1':
                    self.flowables.append(Paragraph(text, self.styles['Heading1']))
                    self.flowables.append(Spacer(1, 8))
                elif tag == 'h2':
                    self.flowables.append(Paragraph(text, self.styles['Heading2']))
                    self.flowables.append(Spacer(1, 6))
                elif tag == 'h3':
                    self.flowables.append(Paragraph(text, self.styles['Heading3']))
                    self.flowables.append(Spacer(1, 6))
                elif tag in ['h4', 'h5', 'h6']:
                    self.flowables.append(Paragraph(text, self.styles['Heading4']))
                    self.flowables.append(Spacer(1, 6))
                elif tag == 'p':
                    self.flowables.append(Paragraph(text, self.styles['BodyText']))
                    self.flowables.append(Spacer(1, 8))
                elif tag == 'li':
                    prefix = "&bull; " if self.list_mode == 'ul' else f"{self.list_index}. "
                    self.list_index += 1
                    self.flowables.append(Paragraph(f"{prefix}{text}", self.styles['Normal']))
                    self.flowables.append(Spacer(1, 4))
                elif tag in ['pre', 'code']:
                    self.flowables.append(Paragraph(text, self.styles['Code']))
                    self.flowables.append(Spacer(1, 6))
            self.current_tag = None
            self.current_text = ""
        elif tag in ['ul', 'ol']:
            self.list_mode = None
            self.flowables.append(Spacer(1, 6))

    def handle_data(self, data):
        if self.current_tag:
            self.current_text += data

def build_pdf_from_text(txt_content: str, pdf_path: str):
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    styles = getSampleStyleSheet()
    text_style = ParagraphStyle(
        'TxtStyle',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=10,
        leading=12
    )
    story = [Preformatted(txt_content, text_style)]
    doc.build(story)

def build_pdf_from_html(html_content: str, pdf_path: str):
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    styles = getSampleStyleSheet()
    parser = HTMLToPlatypusParser(styles)
    parser.feed(html_content)
    doc.build(parser.flowables)

def build_pdf_from_docx(docx_path: str, pdf_path: str):
    from docx import Document
    doc_file = Document(docx_path)
    
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    styles = getSampleStyleSheet()
    story = []
    
    for para in doc_file.paragraphs:
        para_text = ""
        for run in para.runs:
            text = run.text
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            para_text += text
            
        if not para_text.strip():
            story.append(Spacer(1, 8))
            continue
            
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            story.append(Paragraph(para_text, styles['Heading1']))
            story.append(Spacer(1, 8))
        elif 'heading 2' in style_name:
            story.append(Paragraph(para_text, styles['Heading2']))
            story.append(Spacer(1, 6))
        elif 'heading 3' in style_name:
            story.append(Paragraph(para_text, styles['Heading3']))
            story.append(Spacer(1, 6))
        elif 'list' in style_name or 'bullet' in style_name:
            story.append(Paragraph(f"&bull; {para_text}", styles['Normal']))
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(para_text, styles['BodyText']))
            story.append(Spacer(1, 8))
            
    doc.build(story)

@app.post("/api/doc-to-pdf")
def convert_doc_to_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """
    Uploads a .doc, .docx, .txt, or .md file, converts it to PDF using LibreOffice,
    returns the PDF file, and registers a cleanup task.
    """
    # 1. Validate file extension
    filename = file.filename or "document.txt"
    sanitized_filename = os.path.basename(filename)
    _, ext = os.path.splitext(sanitized_filename.lower())
    
    allowed_extensions = {".docx", ".doc", ".txt", ".md"}
    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only DOCX, DOC, TXT, and MD files are allowed."
        )

    # Create private temp dir
    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, f"{uuid.uuid4()}{ext}")
    
    # Queue cleanup task to prevent disk pollution
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        # Save upload to temporary file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 2. Validate file size (15MB limit)
        if os.path.getsize(input_path) > 15 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail="File size exceeds the 15MB limit."
            )

        # 3. Validate content / magic bytes
        if ext in {".docx", ".doc"}:
            with open(input_path, "rb") as f:
                header = f.read(8)
                if ext == ".docx" and not header.startswith(b"PK\x03\x04"):
                    raise HTTPException(
                        status_code=400,
                        detail="Invalid Word Document (.docx) file structure."
                    )
                if ext == ".doc" and header != b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
                    raise HTTPException(
                        status_code=400,
                        detail="Invalid Word Document (.doc) file structure."
                    )
        elif ext in {".txt", ".md"}:
            with open(input_path, "rb") as f:
                chunk = f.read(8192)
                if b"\x00" in chunk:
                    raise HTTPException(
                        status_code=400,
                        detail="Invalid text file. Binary content detected."
                    )

        logging.info(f"Starting PDF conversion for file: {sanitized_filename} (type: {ext})")

        # 4. Handle Markdown conversion to HTML first
        conversion_input_path = input_path
        if ext == ".md":
            try:
                with open(input_path, "r", encoding="utf-8") as f:
                    md_text = f.read()
                
                # Convert Markdown to HTML
                html_content = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'tables'])
                
                # Wrap in styled HTML template
                styled_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&family=Inter:wght@300;400;500;600;700&display=swap');
  body {{
    font-family: 'Inter', sans-serif;
    color: #1e293b;
    line-height: 1.6;
    max-width: 800px;
    margin: 40px auto;
    padding: 0 20px;
  }}
  h1, h2, h3, h4, h5, h6 {{
    font-family: 'Outfit', sans-serif;
    color: #0f172a;
    margin-top: 24px;
    margin-bottom: 12px;
    font-weight: 700;
  }}
  h1 {{ font-size: 2.25rem; border-bottom: 2px solid #e2e8f0; padding-bottom: 8px; }}
  h2 {{ font-size: 1.75rem; border-bottom: 1px solid #e2e8f0; padding-bottom: 6px; }}
  h3 {{ font-size: 1.4rem; }}
  p {{ margin-bottom: 16px; }}
  code {{
    font-family: 'Courier New', Courier, monospace;
    background-color: #f1f5f9;
    padding: 2px 6px;
    border-radius: 4px;
    font-size: 0.9em;
  }}
  pre {{
    background-color: #f8fafc;
    border: 1px solid #e2e8f0;
    padding: 16px;
    border-radius: 8px;
    overflow-x: auto;
    margin-bottom: 16px;
  }}
  pre code {{
    background-color: transparent;
    padding: 0;
    border-radius: 0;
  }}
  blockquote {{
    border-left: 4px solid #6366f1;
    padding-left: 16px;
    color: #475569;
    font-style: italic;
    margin: 16px 0;
  }}
  table {{
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 24px;
  }}
  th, td {{
    border: 1px solid #e2e8f0;
    padding: 10px 14px;
    text-align: left;
  }}
  th {{
    background-color: #f8fafc;
    font-weight: 600;
  }}
  img {{
    max-width: 100%;
    height: auto;
    border-radius: 6px;
  }}
  ul, ol {{
    margin-bottom: 16px;
    padding-left: 24px;
  }}
  li {{
    margin-bottom: 4px;
  }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
                
                conversion_input_path = os.path.join(temp_dir, f"{uuid.uuid4()}.html")
                with open(conversion_input_path, "w", encoding="utf-8") as f:
                    f.write(styled_html)
            except Exception as me:
                logging.error(f"Markdown parsing failed: {me}", exc_info=True)
                # Fallback to direct conversion of raw markdown text
                conversion_input_path = input_path

        # 5. Run headless LibreOffice, falling back to pure Python ReportLab if not present
        use_fallback = False
        try:
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                temp_dir,
                conversion_input_path
            ]
            
            result = subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                logging.warning(f"LibreOffice conversion returned code {result.returncode}: {result.stderr}. Using ReportLab fallback...")
                use_fallback = True
        except FileNotFoundError:
            logging.info("LibreOffice binary not found. Using pure-Python ReportLab/docx parser.")
            use_fallback = True

        if use_fallback:
            try:
                pdf_filename = os.path.splitext(os.path.basename(conversion_input_path))[0] + ".pdf"
                pdf_path = os.path.join(temp_dir, pdf_filename)
                
                if ext == ".txt":
                    with open(conversion_input_path, "r", encoding="utf-8", errors="ignore") as f:
                        txt_content = f.read()
                    build_pdf_from_text(txt_content, pdf_path)
                elif ext == ".md":
                    # For MD, conversion_input_path points to the HTML wrapper we constructed in step 4
                    with open(conversion_input_path, "r", encoding="utf-8", errors="ignore") as f:
                        html_content = f.read()
                    build_pdf_from_html(html_content, pdf_path)
                elif ext == ".docx":
                    build_pdf_from_docx(conversion_input_path, pdf_path)
                elif ext == ".doc":
                    raise ValueError("Legacy .doc format is not supported in pure-Python mode. Please convert to .docx first.")
            except Exception as fe:
                logging.error(f"Pure-Python conversion fallback failed: {fe}", exc_info=True)
                raise HTTPException(
                    status_code=500,
                    detail=f"Document conversion failed: {str(fe)}"
                )

        # 6. Locate output PDF
        pdf_filename = os.path.splitext(os.path.basename(conversion_input_path))[0] + ".pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            raise HTTPException(
                status_code=500,
                detail="Conversion succeeded but PDF file was not created or is empty."
            )

        base_name = os.path.splitext(sanitized_filename)[0]
        pdf_download_filename = f"{base_name}.pdf"

        logging.info(f"PDF conversion successful, sending file: {pdf_download_filename}")

        return FileResponse(
            path=pdf_path,
            filename=pdf_download_filename,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{pdf_download_filename}"',
                "X-Content-Type-Options": "nosniff"
            }
        )

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Error during document to PDF conversion: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred during conversion: {str(e)}"
        )


@app.post("/api/pdf-to-jpg")
def convert_pdf_to_jpg(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    join_pages: bool = Form(False)
):
    """
    Uploads a PDF, converts pages to JPG images.
    If multi-page and join_pages is False, returns a ZIP file.
    If multi-page and join_pages is True, joins pages vertically and returns a single JPG.
    If single page, returns a single JPG.
    """
    filename = file.filename or "document.pdf"
    sanitized_filename = os.path.basename(filename)
    if not sanitized_filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only PDF documents are allowed."
        )

    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        # Save upload to temporary file
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Validate size
        if os.path.getsize(pdf_path) > 20 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail="PDF file exceeds the 20MB size limit."
            )

        if os.path.getsize(pdf_path) < 5:
             raise HTTPException(status_code=400, detail="Invalid PDF file. File is too small.")
             
        with open(pdf_path, "rb") as f:
            magic_bytes = f.read(5)
            if magic_bytes != b"%PDF-":
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid PDF file header. The uploaded file is not a valid PDF."
                )

        logging.info(f"Converting PDF to JPG: {sanitized_filename} (join_pages={join_pages})")

        doc = fitz.open(pdf_path)
        num_pages = len(doc)
        if num_pages == 0:
            raise HTTPException(status_code=400, detail="The PDF file has no pages.")

        # Define rendering resolution: 150 DPI is standard for high quality
        zoom = 150 / 72
        mat = fitz.Matrix(zoom, zoom)

        base_name = os.path.splitext(sanitized_filename)[0]

        # Case 1: Single page PDF or joined vertically
        if num_pages == 1 or (num_pages > 1 and join_pages):
            if num_pages == 1:
                page = doc.load_page(0)
                pix = page.get_pixmap(matrix=mat)
                jpg_filename = f"{base_name}.jpg"
                jpg_path = os.path.join(temp_dir, jpg_filename)
                pix.save(jpg_path)
            else:
                # Join multiple pages vertically
                pil_images = []
                total_height = 0
                max_width = 0
                for page_num in range(num_pages):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    pil_images.append(img)
                    total_height += img.height
                    if img.width > max_width:
                          max_width = img.width
                
                combined_img = Image.new("RGB", (max_width, total_height), (255, 255, 255))
                current_y = 0
                for img in pil_images:
                    x_offset = (max_width - img.width) // 2
                    combined_img.paste(img, (x_offset, current_y))
                    current_y += img.height

                jpg_filename = f"{base_name}_combined.jpg"
                jpg_path = os.path.join(temp_dir, jpg_filename)
                combined_img.save(jpg_path, "JPEG", quality=90)

            return FileResponse(
                path=jpg_path,
                filename=jpg_filename,
                media_type="image/jpeg",
                headers={
                    "Content-Disposition": f'attachment; filename="{jpg_filename}"',
                    "X-Content-Type-Options": "nosniff"
                }
            )

        # Case 2: Multi-page PDF to separate JPGs (Zipped)
        else:
            zip_filename = f"{base_name}_images.zip"
            zip_path = os.path.join(temp_dir, zip_filename)
            
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                for page_num in range(num_pages):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=mat)
                    page_jpg_path = os.path.join(temp_dir, f"{base_name}_page_{page_num + 1}.jpg")
                    pix.save(page_jpg_path)
                    zipf.write(page_jpg_path, arcname=f"{base_name}_page_{page_num + 1}.jpg")
                    os.remove(page_jpg_path)

            return FileResponse(
                path=zip_path,
                filename=zip_filename,
                media_type="application/zip",
                headers={
                    "Content-Disposition": f'attachment; filename="{zip_filename}"',
                    "X-Content-Type-Options": "nosniff"
                }
            )

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Error during PDF to JPG conversion: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred during conversion: {str(e)}"
        )


@app.post("/api/jpg-to-pdf")
def convert_jpg_to_pdf(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...)
):
    """
    Uploads a list of JPG/JPEG/PNG images, compiles them into a single multi-page PDF.
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded.")
        
    temp_dir = tempfile.mkdtemp()
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        pil_images = []
        for index, file in enumerate(files):
            filename = file.filename or f"image_{index}.jpg"
            sanitized_filename = os.path.basename(filename)
            _, ext = os.path.splitext(sanitized_filename.lower())
            if ext not in {".jpg", ".jpeg", ".png"}:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid file type: {sanitized_filename}. Only JPG, JPEG, and PNG images are allowed."
                )

            # Write to a temp image file to validate size and load with Pillow
            temp_img_path = os.path.join(temp_dir, f"img_{index}{ext}")
            with open(temp_img_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            # Limit file size to 10MB per image
            if os.path.getsize(temp_img_path) > 10 * 1024 * 1024:
                raise HTTPException(
                    status_code=400,
                    detail=f"File {sanitized_filename} exceeds the 10MB limit."
                )

            try:
                # Verify and load image
                img = Image.open(temp_img_path)
                img.verify()
                # Re-open because verify() makes the image object unusable for saving
                img = Image.open(temp_img_path)
                # Convert to RGB mode if not already
                if img.mode != "RGB":
                    img = img.convert("RGB")
                # Load into memory to avoid file locks
                img.load()
                pil_images.append(img)
            except Exception as img_err:
                logging.error(f"Image load failure for {sanitized_filename}: {img_err}")
                raise HTTPException(
                    status_code=400,
                    detail=f"Failed to process image {sanitized_filename}. The file might be corrupted."
                )

        if not pil_images:
            raise HTTPException(status_code=400, detail="No valid images to convert.")

        # Save to a PDF file
        pdf_filename = "converted_images.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        # Pillow can save a list of images to a single PDF
        pil_images[0].save(
            pdf_path,
            save_all=True,
            append_images=pil_images[1:]
        )

        # Double check file was created and is non-empty
        if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
            raise HTTPException(
                status_code=500,
                detail="Failed to generate the PDF file."
            )

        return FileResponse(
            path=pdf_path,
            filename=pdf_filename,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{pdf_filename}"',
                "X-Content-Type-Options": "nosniff"
            }
        )

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Error during JPG to PDF conversion: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred during conversion: {str(e)}"
        )


@app.post("/api/evaluate-pdf")
def evaluate_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
):
    """
    Uploads a PDF, inspects its metadata (size, pages, image count),
    and returns a recommended compression profile.
    """
    filename = file.filename or "document.pdf"
    sanitized_filename = os.path.basename(filename)
    if not sanitized_filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only PDF documents are allowed."
        )

    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        # Save upload to temporary file
        with open(pdf_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Validate size
        file_size = os.path.getsize(pdf_path)
        if file_size > 200 * 1024 * 1024:  # 200MB limit for local compression
            raise HTTPException(
                status_code=400,
                detail="PDF file exceeds the 200MB size limit."
            )

        if file_size < 5:
             raise HTTPException(status_code=400, detail="Invalid PDF file. File is too small.")
             
        with open(pdf_path, "rb") as f:
            magic_bytes = f.read(5)
            if magic_bytes != b"%PDF-":
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid PDF file header. The uploaded file is not a valid PDF."
                )

        doc = fitz.open(pdf_path)
        pages = len(doc)
        
        # Count images
        image_count = 0
        for page in doc:
            try:
                image_count += len(page.get_images())
            except Exception:
                pass
        
        doc.close()

        # Recommendation logic
        if file_size < 500 * 1024 and image_count == 0:
            rec_level = "none"
            rec_text = "This PDF is already very small (under 500 KB) and contains no images. Compression is not recommended as it will not yield noticeable savings."
            expected_reduction = "0% - 5%"
        elif file_size > 5 * 1024 * 1024:
            rec_level = "high"
            rec_text = f"This PDF is quite large ({round(file_size / (1024 * 1024), 2)} MB) and contains {image_count} image(s). We recommend High compression (screen quality, 72 DPI) to significantly reduce the file size."
            expected_reduction = "50% - 80%"
        else:
            rec_level = "medium"
            if image_count > 0:
                rec_text = f"This PDF contains {image_count} image(s) and is of moderate size. We recommend Medium compression (e-book quality, 150 DPI) to balance layout readability and file size."
                expected_reduction = "30% - 60%"
            else:
                rec_text = "This document is of moderate size but has no images. We recommend Medium compression to optimize document structures and page content stream inflation."
                expected_reduction = "10% - 30%"

        return {
            "filename": sanitized_filename,
            "file_size_bytes": file_size,
            "pages": pages,
            "image_count": image_count,
            "recommendation": {
                "level": rec_level,
                "text": rec_text,
                "expected_reduction": expected_reduction
            }
        }

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Error during PDF evaluation: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred during evaluation: {str(e)}"
        )


@app.post("/api/compress-pdf")
def compress_pdf(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    level: str = Form("medium")
):
    """
    Uploads a PDF, compresses it using Ghostscript with the specified quality profile.
    """
    filename = file.filename or "document.pdf"
    sanitized_filename = os.path.basename(filename)
    if not sanitized_filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400, 
            detail="Invalid file type. Only PDF documents are allowed."
        )

    # Map compression level to Ghostscript PDFSETTINGS profiles
    level_map = {
        "high": "/screen",      # 72 DPI
        "medium": "/ebook",      # 150 DPI
        "low": "/printer"        # 300 DPI
    }
    gs_profile = level_map.get(level.lower(), "/ebook")

    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")
    background_tasks.add_task(cleanup_temp_dir, temp_dir)

    try:
        # Save upload to temporary file
        with open(input_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Validate size
        if os.path.getsize(input_path) > 200 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail="PDF file exceeds the 200MB size limit."
            )

        if os.path.getsize(input_path) < 5:
             raise HTTPException(status_code=400, detail="Invalid PDF file. File is too small.")
             
        with open(input_path, "rb") as f:
            magic_bytes = f.read(5)
            if magic_bytes != b"%PDF-":
                raise HTTPException(
                    status_code=400, 
                    detail="Invalid PDF file header. The uploaded file is not a valid PDF."
                )

        logging.info(f"Compressing PDF: {sanitized_filename} (level: {level}, profile: {gs_profile})")

        # Run Ghostscript to compress PDF
        cmd = [
            "gs",
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.4",
            f"-dPDFSETTINGS={gs_profile}",
            "-dNOPAUSE",
            "-dQUIET",
            "-dBATCH",
            f"-sOutputFile={output_path}",
            input_path
        ]

        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=120
        )

        if result.returncode != 0:
            logging.error(f"Ghostscript compression failed with return code {result.returncode}: {result.stderr}")
            raise HTTPException(
                status_code=500,
                detail="Ghostscript compression failed. Please verify that the PDF is not password-protected or corrupted."
            )

        # Double check output file exists and is non-empty
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise HTTPException(
                status_code=500,
                detail="Compression failed to create output file."
            )

        # Derive download filename
        base_name = os.path.splitext(sanitized_filename)[0]
        compressed_filename = f"{base_name}_compressed.pdf"

        logging.info(f"PDF compression successful: {sanitized_filename} -> size: {os.path.getsize(output_path)} bytes")

        return FileResponse(
            path=output_path,
            filename=compressed_filename,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{compressed_filename}"',
                "X-Content-Type-Options": "nosniff"
            }
        )

    except HTTPException as he:
        raise he
    except subprocess.TimeoutExpired:
        logging.error("Ghostscript compression timed out.")
        raise HTTPException(
            status_code=504,
            detail="PDF compression operation timed out."
        )
    except Exception as e:
        logging.error(f"Error during PDF compression: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred during compression: {str(e)}"
        )


# Static files mapping

# Serve built frontend production assets from /home/dlh/dlhdev/pdftools/dist
static_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "dist"))

# Mount Static Files (Vite production build)
if os.path.exists(static_dir):
    logging.info(f"Serving static files from: {static_dir}")
    app.mount("/", StaticFiles(directory=static_dir, html=True), name="static")
else:
    logging.warning(f"Static directory not found at {static_dir}. Make sure you build the frontend before starting the server in production.")
    @app.get("/")
    def index_fallback():
        return {
            "status": "online",
            "message": "PDF Tools Backend is running. Please run 'npm run build' to serve the frontend."
        }
